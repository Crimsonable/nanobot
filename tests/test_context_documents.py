"""Tests for context builder media handling.

The ContextBuilder._build_user_content method should ONLY handle images.
Document text extraction is the responsibility of the processing layer
(AgentLoop._process_message and _drain_pending).
"""

from __future__ import annotations

from pathlib import Path

from nanobot.agent.context import ContextBuilder
from nanobot.utils.document import extract_documents


def _make_builder(tmp_path: Path) -> ContextBuilder:
    """Create a minimal ContextBuilder for testing."""
    return ContextBuilder(workspace=tmp_path, timezone="UTC")


def test_build_user_content_with_no_media_returns_string(tmp_path: Path) -> None:
    builder = _make_builder(tmp_path)
    result = builder._build_user_content("hello", None)
    assert result == "hello"


def test_build_user_content_with_image_returns_list(tmp_path: Path) -> None:
    """Normalized image refs should produce multimodal content blocks."""
    builder = _make_builder(tmp_path)
    result = builder._build_user_content("describe this", ["image_url:https://files.example/test.png"])
    assert isinstance(result, list)
    types = [b["type"] for b in result]
    assert "image_url" in types
    assert "text" in types


def test_build_user_content_ignores_untyped_media_refs(tmp_path: Path) -> None:
    """Untyped refs should be ignored because extract_documents must normalize them first."""
    builder = _make_builder(tmp_path)
    result = builder._build_user_content("summarize", ["https://files.example/notes.txt"])
    assert result == "summarize"


def test_build_user_content_mixed_typed_and_untyped_refs(tmp_path: Path) -> None:
    """Only normalized visual refs should be included."""
    builder = _make_builder(tmp_path)

    result = builder._build_user_content(
        "analyze",
        ["image_url:https://files.example/chart.png", "https://files.example/report.txt"],
    )
    assert isinstance(result, list)
    assert any(b["type"] == "image_url" for b in result)
    assert all(b.get("type") != "video_url" for b in result if isinstance(b, dict))


def test_build_user_content_with_video_returns_video_block(tmp_path: Path) -> None:
    builder = _make_builder(tmp_path)

    result = builder._build_user_content("watch this", ["video_url:https://files.example/demo.mp4"])

    assert isinstance(result, list)
    assert any(b.get("type") == "video_url" for b in result if isinstance(b, dict))


# ---------------------------------------------------------------------------
# Bug detection: extract_documents must be called BEFORE _build_user_content
# to prevent document media from being silently dropped.
# This simulates the _drain_pending code path.
# ---------------------------------------------------------------------------

def test_drain_pending_path_preserves_document_reference(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """Pending follow-ups should retain document URLs in text form."""
    from docx import Document
    import nanobot.utils.document as _doc

    doc = Document()
    doc.add_paragraph("Quarterly revenue is $5M")
    docx_path = tmp_path / "report.docx"
    doc.save(docx_path)

    monkeypatch.setattr(
        _doc,
        "_upload_local_attachment_ref",
        lambda path, metadata=None: f"https://files.example/{path.name}",
    )

    content = "summarize"
    media = [str(docx_path)]

    # Step 1: extract_documents separates visual media from ordinary files.
    new_content, image_only = extract_documents(content, media)

    # Step 2: _build_user_content handles only visual media (none left here).
    builder = _make_builder(tmp_path)
    result = builder._build_user_content(new_content, image_only if image_only else None)

    assert isinstance(result, str)
    assert "summarize" in result
    assert "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in result
    assert "https://files.example/report.docx" in result


def test_drain_pending_path_without_extract_loses_document(tmp_path: Path) -> None:
    """Without extract_documents, direct document attachments are still ignored."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Secret data in document")
    docx_path = tmp_path / "report.docx"
    doc.save(docx_path)

    builder = _make_builder(tmp_path)

    # Raw document refs are not handled here; extract_documents must normalize them first.
    result = builder._build_user_content("summarize", [str(docx_path)])

    assert result == "summarize"
