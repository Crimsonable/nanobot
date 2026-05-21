"""
智成安监 - 报告生成器
输入 JSON，渲染 HTML 字符串，并使用 Playwright 导出 PDF。
"""

from __future__ import annotations

import argparse
from io import BytesIO
import json
import mimetypes
import os
from datetime import datetime
from pathlib import Path
from typing import Any
import httpx
from urllib.request import urlopen
from urllib.error import URLError

from jinja2 import Template

try:
    from playwright.sync_api import sync_playwright  # type: ignore

    PLAYWRIGHT_AVAILABLE = True
except ImportError:
    PLAYWRIGHT_AVAILABLE = False

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt, Cm, RGBColor
    from docx.oxml import OxmlElement

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def _set_run_font(run, font_name="Microsoft YaHei", font_size=11):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(font_size)


def _set_para_font(para, font_name="Microsoft YaHei", font_size=11):
    for run in para.runs:
        _set_run_font(run, font_name, font_size)


def _set_cell_font(cell, font_size=11):
    for para in cell.paragraphs:
        for run in para.runs:
            _set_run_font(run, "Microsoft YaHei", font_size)


def _set_heading_font(style, font_name="Microsoft YaHei", font_size=14):
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    style.font.size = Pt(font_size)


def _get_image_source(photo_ref: str | None) -> str | BytesIO | None:
    """将图片路径（本地或URL）转为 python-docx 可接受的路径或内存流"""
    if not photo_ref:
        return None
    if os.path.exists(photo_ref):
        return photo_ref
    if photo_ref.startswith(("http://", "https://")):
        try:
            with urlopen(photo_ref, timeout=10) as response:
                data = response.read()
            return BytesIO(data)
        except (URLError, OSError):
            return None
    return None


DEFAULT_TEMPLATE = """<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>成都院智成安监检查反馈意见</title>
<style>
:root{
  --bg:#f5f7fa;
  --card:#ffffff;
  --border:#e2e8f0;
  --text:#1a202c;
  --text2:#64748b;
  --red:#dc2626;
  --orange:#ea580c;
  --blue:#2563eb;
  --green:#16a34a;
  --red-bg:#fef2f2;
  --red-bd:#fecaca;
  --orange-bg:#fff7ed;
  --orange-bd:#fed7aa;
  --blue-bg:#eff6ff;
  --blue-bd:#bfdbfe;
  --green-bg:#f0fdf4;
  --green-bd:#bbf7d0;
}
*{box-sizing:border-box;margin:0;padding:0;}
html,body{
  background:var(--bg);
  font-family:"PingFang SC","Microsoft YaHei","Helvetica Neue",Arial,sans-serif;
  color:var(--text);
  line-height:1.7;
  font-size:14px;
}
.container{max-width:900px;margin:0 auto;padding:24px 16px;}
.report-title{text-align:center;font-size:1.5rem;font-weight:700;margin-bottom:24px;}
.card{
  background:var(--card);
  border-radius:12px;
  padding:20px 24px;
  margin-bottom:16px;
  border:1px solid var(--border);
}
.section-header{
  font-size:1rem;font-weight:700;margin-bottom:16px;padding-bottom:10px;
  border-bottom:1px solid var(--border);display:flex;align-items:center;gap:8px;
}
.section-header .num{
  display:inline-flex;align-items:center;justify-content:center;width:24px;height:24px;
  border-radius:50%;font-size:0.75rem;font-weight:700;color:#fff;flex-shrink:0;
}
.info-grid{display:grid;grid-template-columns:repeat(2,1fr);gap:12px 32px;}
.info-item{display:flex;font-size:0.9rem;}
.info-label{color:var(--text2);min-width:100px;flex-shrink:0;font-weight:500;}
.scene-keywords{display:flex;flex-wrap:wrap;gap:8px;}
.keyword-tag{
  padding:4px 12px;background:var(--blue-bg);border:1px solid var(--blue-bd);
  border-radius:6px;font-size:0.8rem;color:var(--blue);font-weight:500;
}
.summary-row{display:flex;gap:16px;margin-bottom:16px;flex-wrap:wrap;}
.summary-box{flex:1;min-width:120px;padding:14px 20px;border-radius:10px;text-align:center;}
.summary-box.red{background:var(--red-bg);border:1px solid var(--red-bd)}
.summary-box.orange{background:var(--orange-bg);border:1px solid var(--orange-bd)}
.summary-box.blue{background:var(--blue-bg);border:1px solid var(--blue-bd)}
.summary-box .num{font-size:1.8rem;font-weight:800;line-height:1}
.summary-box .label{margin-top:4px;color:var(--text2);font-size:0.8rem}
.table-wrapper{overflow-x:auto;}
table{width:100%;border-collapse:collapse;font-size:0.85rem;}
th{
  background:#f8fafc;padding:10px 12px;text-align:left;font-weight:600;
  border:1px solid var(--border);white-space:nowrap;
}
td{padding:10px 12px;border:1px solid var(--border);vertical-align:top;}
.tr-highlight td{background:#fff5f5;}
.tr-warning td{background:#fff7ed;}
.tr-suggest td{background:#eff6ff;}
.hazard-desc{font-weight:500;}
.clause-item{margin-bottom:4px;font-size:0.8rem;line-height:1.5;}
.clause-item strong{color:var(--blue);font-weight:600;}
.photo-cell{width:120px;}
.photo-cell img{max-width:100px;max-height:80px;border-radius:6px;border:1px solid var(--border);}
.remark-cell{color:var(--text2);font-size:0.8rem;}
.reminder-content{
  padding:12px 16px;background:#f8fafc;border-radius:8px;border-left:4px solid var(--orange);
  font-size:0.9rem;line-height:1.8;
}
.footer{
  text-align:center;padding:20px;color:var(--text2);font-size:0.8rem;
  border-top:1px solid var(--border);margin-top:16px;
}
.footer .company{font-weight:600;color:var(--blue);margin-bottom:4px;}
</style>
</head>
<body>
<div class="container">
<div class="report-title">成都院智成安监检查反馈意见</div>

<div class="card">
  <div class="section-header"><span class="num" style="background:var(--blue)">1</span>基本信息</div>
  <div class="info-grid">
    <div class="info-item"><span class="info-label">提交人：</span><span>{{ submitter }}</span></div>
    <div class="info-item"><span class="info-label">检查时间：</span><span>{{ inspect_time }}</span></div>
    <div class="info-item"><span class="info-label">报告编号：</span><span>{{ report_id }}</span></div>
  </div>
</div>

<div class="card">
  <div class="section-header"><span class="num" style="background:var(--green)">2</span>场景认定</div>
  <div class="scene-keywords">
    {% for keyword in scene_keywords %}
    <span class="keyword-tag">{{ keyword }}</span>
    {% endfor %}
  </div>
</div>

<div class="card">
  <div class="section-header"><span class="num" style="background:var(--orange)">3</span>现场事故隐患总体评价</div>
  <div class="summary-row">
    <div class="summary-box red"><div class="num">{{ severe_count }}</div><div class="label">重大事故隐患</div></div>
    <div class="summary-box orange"><div class="num">{{ general_count }}</div><div class="label">一般事故隐患</div></div>
    <div class="summary-box blue"><div class="num">{{ improvement_items|length }}</div><div class="label">改进提升事项</div></div>
  </div>
  <p style="margin-top:12px;font-size:0.9rem;color:var(--text2);">{{ overall_evaluation }}</p>
</div>

<div class="card">
  <div class="section-header"><span class="num" style="background:var(--red)">4</span>现场事故隐患清单</div>
  <div style="margin-bottom:24px">
    <div class="section-header" style="font-size:0.95rem;margin-bottom:12px;padding-bottom:8px">（一）重大事故隐患</div>
    <div class="table-wrapper"><table><thead><tr>
      <th style="width:50px">序号</th><th style="width:120px">现场照片</th><th>存在的重大事故隐患</th><th>判定依据</th><th style="width:120px">备注</th>
    </tr></thead><tbody>
      {% for hazard in severe_hazards %}
      <tr class="tr-highlight">
        <td>{{ hazard.index }}</td>
        {% if hazard.rowspan > 0 %}<td class="photo-cell" rowspan="{{ hazard.rowspan }}">{% if hazard.photo_url %}<img src="{{ hazard.photo_url }}" alt="现场照片">{% endif %}</td>{% endif %}
        <td class="hazard-desc">{{ hazard.description }}</td>
        <td>
          <div class="clause-item"><strong>规范：</strong>{{ hazard.standard_name }} {{ hazard.clause_no }}</div>
          <div class="clause-item"><strong>条款原文：</strong>{{ hazard.clause_text }}</div>
        </td>
        <td class="remark-cell">{{ hazard.remark }}</td>
      </tr>
      {% endfor %}
    </tbody></table></div>
  </div>
  <div>
    <div class="section-header" style="font-size:0.95rem;margin-bottom:12px;padding-bottom:8px">（二）一般事故隐患</div>
    <div class="table-wrapper"><table><thead><tr>
      <th style="width:50px">序号</th><th style="width:120px">现场照片</th><th>存在的一般事故隐患</th><th>判定依据</th><th style="width:120px">备注</th>
    </tr></thead><tbody>
      {% for hazard in general_hazards %}
      <tr class="tr-warning">
        <td>{{ hazard.index }}</td>
        {% if hazard.rowspan > 0 %}<td class="photo-cell" rowspan="{{ hazard.rowspan }}">{% if hazard.photo_url %}<img src="{{ hazard.photo_url }}" alt="现场照片">{% endif %}</td>{% endif %}
        <td class="hazard-desc">{{ hazard.description }}</td>
        <td>
          <div class="clause-item"><strong>规范：</strong>{{ hazard.standard_name }} {{ hazard.clause_no }}</div>
          <div class="clause-item"><strong>条款原文：</strong>{{ hazard.clause_text }}</div>
        </td>
        <td class="remark-cell">{{ hazard.remark }}</td>
      </tr>
      {% endfor %}
    </tbody></table></div>
  </div>
</div>

<div class="card">
  <div class="section-header"><span class="num" style="background:var(--green)">5</span>现场改进提升事项清单</div>
  <div class="table-wrapper"><table><thead><tr>
    <th style="width:50px">序号</th><th style="width:120px">现场照片</th><th>存在的改进提升事项</th><th>判定依据</th><th style="width:120px">备注</th>
  </tr></thead><tbody>
    {% for item in improvement_items %}
    <tr class="tr-suggest">
      <td>{{ item.index }}</td>
      {% if item.rowspan > 0 %}<td class="photo-cell" rowspan="{{ item.rowspan }}">{% if item.photo_url %}<img src="{{ item.photo_url }}" alt="现场照片">{% endif %}</td>{% endif %}
      <td class="hazard-desc">{{ item.description }}</td>
      <td>
        <div class="clause-item"><strong>规范：</strong>{{ item.standard_name }} {{ item.clause_no }}</div>
        <div class="clause-item"><strong>条款原文：</strong>{{ item.clause_text }}</div>
      </td>
      <td class="remark-cell">{{ item.remark }}</td>
    </tr>
      {% endfor %}
  </tbody></table></div>
</div>

<div class="card">
  <div class="section-header"><span class="num" style="background:var(--orange)">6</span>下一步安全生产工作提醒</div>
  <div class="reminder-content">
    {% for reminder in next_steps_reminder %}
    <p style="margin-bottom:8px">{{ loop.index }}. {{ reminder }}</p>
    {% endfor %}
  </div>
</div>

<div class="footer">
  <div class="company">中国电建成都院智成安监系统</div>
  <div style="margin-top:4px">生成时间：{{ generated_time }}</div>
</div>
</div>
</body>
</html>"""

DEFAULT_UPLOAD_TARGET = "http://192.168.48.104:30080/internal/attachments/upload"
DEFAULT_FRONTEND_ID = "web-main"
DEFAULT_USER_ID = ""
SCRIPT_DIR = Path(__file__).resolve().parent
SCRIPT_CONFIG_PATH = SCRIPT_DIR / "report_generator.config.json"

REQUIRED_TOP_LEVEL_FIELDS = [
    "submitter",
    "inspect_time",
    "report_id",
    "image_url_map",
    "scene_keywords",
    "severe_count",
    "general_count",
    "overall_evaluation",
    "severe_hazards",
    "general_hazards",
    "improvement_items",
    "next_steps_reminder",
]

REQUIRED_ITEM_FIELDS = [
    "index",
    "photo_name",
    "description",
    "standard_name",
    "clause_no",
    "clause_text",
    "remark",
]


def _guess_content_type(path: Path) -> str:
    content_type, _ = mimetypes.guess_type(str(path))
    return content_type or "application/octet-stream"


def _load_script_config(path: Path = SCRIPT_CONFIG_PATH) -> dict[str, Any]:
    if not path.is_file():
        return {}
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except Exception as exc:
        raise RuntimeError(f"读取脚本配置失败: {path}: {exc}") from exc
    if not isinstance(data, dict):
        raise RuntimeError(f"脚本配置必须是 JSON 对象: {path}")
    return data


def _config_env_or_default(
    config: dict[str, Any],
    config_key: str,
    default: str,
) -> str:
    config_value = str(config.get(config_key) or "").strip()
    if config_value:
        return config_value
    return default


class ObjectStorageUploader:
    def __init__(self, target: str, *, frontend_id: str, user_id: str) -> None:
        self.target = target.rstrip("/")
        self.frontend_id = frontend_id
        self.user_id = user_id

    def upload_file(self, source_path: Path, object_path: Path) -> str:
        return self.upload_bytes(
            filename=object_path.name,
            content=source_path.read_bytes(),
            content_type=_guess_content_type(source_path),
            object_path=object_path,
        )

    def upload_bytes(
        self,
        *,
        filename: str,
        content: bytes,
        content_type: str,
        object_path: Path | None = None,
    ) -> str:
        if self.target.startswith(("http://", "https://")):
            return self._upload_via_http_post(content, filename, content_type)
        if object_path is None:
            object_path = Path(filename)
        return self._upload_bytes_to_local_dir(content, object_path)

    def _upload_via_http_post(self, content: bytes, upload_name: str, content_type: str) -> str:
        data = {"frontend_id": self.frontend_id}
        if self.user_id:
            data["user_id"] = self.user_id
        files = {
            "file": (upload_name, BytesIO(content), content_type),
        }
        try:
            with httpx.Client(timeout=30) as client:
                response = client.post(self.target, data=data, files=files)
                response.raise_for_status()
                payload = response.json()
        except Exception as exc:
            raise RuntimeError(f"上传失败: {exc}") from exc
        uploaded_url = str(payload.get("url") or "").strip()
        if not uploaded_url:
            raise RuntimeError("上传接口未返回 url")
        return uploaded_url

    def _upload_bytes_to_local_dir(self, content: bytes, object_path: Path) -> str:
        target_path = Path(self.target) / object_path
        target_path.parent.mkdir(parents=True, exist_ok=True)
        target_path.write_bytes(content)
        return str(target_path)


def _upload_report_bundle(
    uploader: ObjectStorageUploader,
    report_outputs: list[tuple[str, bytes, str]],
    input_json_bytes: bytes,
    base_name: str,
) -> dict[str, str]:
    uploaded_outputs: dict[str, str] = {}
    for filename, content, content_type in report_outputs:
        object_path = Path(filename)
        uploaded_outputs[f"{Path(filename).suffix.lstrip('.')}_url"] = uploader.upload_bytes(
            filename=filename,
            content=content,
            content_type=content_type,
            object_path=object_path,
        )

    json_object_path = Path(f"{base_name}.json")
    uploaded_outputs["json_url"] = uploader.upload_bytes(
        filename=f"{base_name}.json",
        content=input_json_bytes,
        content_type="application/json; charset=utf-8",
        object_path=json_object_path,
    )
    return uploaded_outputs


class ReportGenerator:
    def __init__(self, template_path: str | None = None) -> None:
        if template_path and os.path.exists(template_path):
            self.template = Template(Path(template_path).read_text(encoding="utf-8"))
        else:
            self.template = Template(DEFAULT_TEMPLATE)

    @staticmethod
    def _with_defaults(report_data: dict[str, Any]) -> dict[str, Any]:
        data = dict(report_data)
        ReportGenerator._validate_report_data(data)
        data.setdefault("generated_time", datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        data["submitter"] = ReportGenerator._normalize_text(data.get("submitter"))
        data["inspect_time"] = ReportGenerator._normalize_text(data.get("inspect_time"))
        data["report_id"] = ReportGenerator._normalize_text(data.get("report_id"))
        data["overall_evaluation"] = ReportGenerator._normalize_text(data.get("overall_evaluation"))
        data["scene_keywords"] = ReportGenerator._normalize_text_list(data.get("scene_keywords"))
        data["next_steps_reminder"] = ReportGenerator._normalize_text_list(
            data.get("next_steps_reminder")
        )
        data["severe_hazards"] = ReportGenerator._resolve_photo_refs(
            data["severe_hazards"], data["image_url_map"]
        )
        data["general_hazards"] = ReportGenerator._resolve_photo_refs(
            data["general_hazards"], data["image_url_map"]
        )
        data["improvement_items"] = ReportGenerator._resolve_photo_refs(
            data["improvement_items"], data["image_url_map"]
        )
        data["severe_hazards"] = ReportGenerator._normalize_items(
            data["severe_hazards"], placeholder_description="无"
        )
        data["general_hazards"] = ReportGenerator._normalize_items(
            data["general_hazards"], placeholder_description="无"
        )
        data["improvement_items"] = ReportGenerator._normalize_items(
            data["improvement_items"], placeholder_description="无"
        )
        return data

    @staticmethod
    def _normalize_text(value: Any, default: str = "无") -> str:
        if not isinstance(value, str):
            return default
        text = value.strip()
        return text or default

    @classmethod
    def _normalize_text_list(cls, values: Any) -> list[str]:
        if not isinstance(values, list):
            return ["无"]
        normalized = [cls._normalize_text(value) for value in values]
        return normalized or ["无"]

    @classmethod
    def _normalize_items(
        cls, items: list[dict[str, Any]], placeholder_description: str = "无"
    ) -> list[dict[str, Any]]:
        if not items:
            return [
                {
                    "index": "-",
                    "photo_name": "",
                    "photo_url": None,
                    "description": placeholder_description,
                    "standard_name": "无",
                    "clause_no": "",
                    "clause_text": "无",
                    "remark": "无",
                }
            ]

        normalized: list[dict[str, Any]] = []
        for item in items:
            row = dict(item)
            row["index"] = row.get("index", "-")
            row["photo_name"] = row.get("photo_name") or ""
            row["photo_url"] = row.get("photo_url")
            row["description"] = cls._normalize_text(row.get("description"))
            row["standard_name"] = cls._normalize_text(row.get("standard_name"))
            row["clause_no"] = cls._normalize_text(row.get("clause_no"))
            row["clause_text"] = cls._normalize_text(row.get("clause_text"))
            row["remark"] = cls._normalize_text(row.get("remark"))
            normalized.append(row)
        return normalized

    @staticmethod
    def _validate_report_data(report_data: dict[str, Any]) -> None:
        errors: list[str] = []
        missing_fields = [
            field for field in REQUIRED_TOP_LEVEL_FIELDS if field not in report_data
        ]
        if missing_fields:
            raise ValueError(
                "输入 JSON 缺少一级字段，共发现以下缺失项：\n"
                + "\n".join(f"  {i + 1}. {field}" for i, field in enumerate(missing_fields))
                + "\n必须完整提供以下一级字段："
                + ", ".join(REQUIRED_TOP_LEVEL_FIELDS)
            )

        image_url_map = report_data["image_url_map"]
        image_url_map_valid = isinstance(image_url_map, dict)
        if not image_url_map_valid:
            errors.append("字段 image_url_map 必须是对象，且值必须是系统实际传入的真实图片 URL")
            image_url_map = {}

        for photo_name, photo_url in image_url_map.items():
            if not isinstance(photo_name, str) or not photo_name.strip():
                errors.append("image_url_map 的 key 必须是非空图片名，例如 img_001.jpg")
            if not isinstance(photo_url, str) or not photo_url.startswith(("http://", "https://")):
                errors.append(
                    f"image_url_map['{photo_name}'] 必须是系统实际传入的真实图片 URL，"
                    "且必须以 http:// 或 https:// 开头"
                )

        list_fields = [
            "scene_keywords",
            "severe_hazards",
            "general_hazards",
            "improvement_items",
            "next_steps_reminder",
        ]
        valid_lists: dict[str, bool] = {}
        for field in list_fields:
            valid_lists[field] = isinstance(report_data[field], list)
            if not valid_lists[field]:
                errors.append(f"字段 {field} 必须是数组")

        if not isinstance(report_data["severe_count"], int):
            errors.append("字段 severe_count 必须是整数")
        if not isinstance(report_data["general_count"], int):
            errors.append("字段 general_count 必须是整数")
        if isinstance(report_data["severe_count"], int) and valid_lists["severe_hazards"]:
            if report_data["severe_count"] != len(report_data["severe_hazards"]):
                errors.append("字段 severe_count 必须与 severe_hazards 数组长度一致")
        if isinstance(report_data["general_count"], int) and valid_lists["general_hazards"]:
            if report_data["general_count"] != len(report_data["general_hazards"]):
                errors.append("字段 general_count 必须与 general_hazards 数组长度一致")

        for field in ("submitter", "inspect_time", "report_id", "overall_evaluation"):
            if not isinstance(report_data[field], str) or not report_data[field].strip():
                errors.append(f"字段 {field} 必须是非空字符串")

        if valid_lists["severe_hazards"]:
            ReportGenerator._validate_items(
                report_data["severe_hazards"], "severe_hazards", image_url_map, errors
            )
        if valid_lists["general_hazards"]:
            ReportGenerator._validate_items(
                report_data["general_hazards"], "general_hazards", image_url_map, errors
            )
        if valid_lists["improvement_items"]:
            ReportGenerator._validate_items(
                report_data["improvement_items"], "improvement_items", image_url_map, errors
            )

        if errors:
            raise ValueError(
                "输入 JSON 校验失败，共发现以下问题：\n"
                + "\n".join(f"  {i + 1}. {error}" for i, error in enumerate(errors))
            )

    @staticmethod
    def _validate_items(
        items: list[dict[str, Any]],
        field_name: str,
        image_url_map: dict[str, str],
        errors: list[str],
    ) -> None:
        for index, item in enumerate(items, start=1):
            if not isinstance(item, dict):
                errors.append(f"{field_name}[{index}] 必须是对象")
                continue
            missing_fields = [field for field in REQUIRED_ITEM_FIELDS if field not in item]
            if missing_fields:
                missing = ", ".join(missing_fields)
                errors.append(
                    f"{field_name}[{index}] 缺少字段：{missing}。"
                    f"每个条目必须包含：{', '.join(REQUIRED_ITEM_FIELDS)}"
                )
                continue

            photo_name = item["photo_name"]
            if not isinstance(photo_name, str) or not photo_name.strip():
                errors.append(f"{field_name}[{index}].photo_name 必须是非空字符串")
            elif photo_name not in image_url_map:
                errors.append(
                    f"{field_name}[{index}].photo_name={photo_name} 未在 image_url_map 中登记。"
                    "image_url_map 必须覆盖报告中实际引用到的全部 photo_name，"
                    "且映射值必须是系统实际传入的真实图片 URL"
                )

            remark = item["remark"]
            if not isinstance(remark, str) or not remark.strip():
                errors.append(f"{field_name}[{index}].remark 必须是非空字符串")

            for required in (
                "description",
                "standard_name",
                "clause_no",
                "clause_text",
            ):
                value = item[required]
                if not isinstance(value, str) or not value.strip():
                    errors.append(f"{field_name}[{index}].{required} 必须是非空字符串")

    @staticmethod
    def _resolve_photo_ref(photo_ref: Any, image_url_map: dict[str, Any]) -> Any:
        if not isinstance(photo_ref, str) or not photo_ref.strip():
            return photo_ref

        ref = photo_ref.strip()
        return image_url_map.get(ref)

    @classmethod
    def _resolve_photo_refs(
        cls, items: list[dict[str, Any]], image_url_map: dict[str, Any]
    ) -> list[dict[str, Any]]:
        resolved: list[dict[str, Any]] = []
        for item in items:
            row = dict(item)
            photo_name = row.get("photo_name")
            row["photo_name"] = photo_name
            row["photo_url"] = cls._resolve_photo_ref(photo_name, image_url_map)
            resolved.append(row)
        return resolved

    @staticmethod
    def _add_rowspan(items: list[dict[str, Any]]) -> list[dict[str, Any]]:
        grouped: dict[str, list[dict[str, Any]]] = {}
        for item in items:
            grouped.setdefault(item.get("photo_name") or "", []).append(item)
        result: list[dict[str, Any]] = []
        for group_items in grouped.values():
            for index, item in enumerate(group_items):
                row = dict(item)
                row["rowspan"] = len(group_items) if index == 0 else 0
                result.append(row)
        return result

    def generate_html(self, report_data: dict[str, Any]) -> str:
        data = self._with_defaults(report_data)
        data["severe_hazards"] = self._add_rowspan(data["severe_hazards"])
        data["general_hazards"] = self._add_rowspan(data["general_hazards"])
        data["improvement_items"] = self._add_rowspan(data["improvement_items"])
        return self.template.render(**data)

    def generate_html_bytes(self, report_data: dict[str, Any]) -> bytes:
        return self.generate_html(report_data).encode("utf-8")

    def html_to_pdf_bytes(self, html: str) -> bytes:
        if not PLAYWRIGHT_AVAILABLE:
            raise ImportError(
                "导出 PDF 需要 playwright，请先执行: pip install playwright && playwright install chromium"
            )
        with sync_playwright() as p:
            browser = p.chromium.launch()
            page = browser.new_page()
            page.set_content(html, wait_until="networkidle")
            pdf_bytes = page.pdf(
                format="A4",
                print_background=True,
                margin={"top": "12mm", "right": "10mm", "bottom": "12mm", "left": "10mm"},
            )
            browser.close()
        return pdf_bytes

    def _build_word_document(self, report_data: dict[str, Any]) -> Any:
        if not DOCX_AVAILABLE:
            raise ImportError("需要安装 python-docx 库: pip install python-docx")
        data = self._with_defaults(report_data)
        doc = Document()

        style = doc.styles["Normal"]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)

        title = doc.add_heading("成都院智成安监检查反馈意见", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if title.runs:
            _set_run_font(title.runs[0], "Microsoft YaHei", 16)

        # 设置所有标题样式字体
        for level in range(1, 7):
            try:
                heading_style = doc.styles[f"Heading {level}"]
                _set_heading_font(heading_style, "Microsoft YaHei", 14 - level * 0.5)
            except Exception:
                pass

        doc.add_heading("1. 基本信息", level=1)
        for run in doc.paragraphs[-1].runs:
            _set_run_font(run, "Microsoft YaHei", 12)
        table = doc.add_table(rows=3, cols=2)
        table.style = "Table Grid"
        for i, (k, v) in enumerate(
            [
                ("提交人", data.get("submitter", "待填写")),
                ("检查时间", data.get("inspect_time", "")),
                ("报告编号", data.get("report_id", "")),
            ]
        ):
            table.cell(i, 0).text = k
            table.cell(i, 1).text = str(v)
            _set_cell_font(table.cell(i, 0))
            _set_cell_font(table.cell(i, 1))

        doc.add_heading("2. 场景认定", level=1)
        for run in doc.paragraphs[-1].runs:
            _set_run_font(run, "Microsoft YaHei", 12)
        if data.get("scene_keywords"):
            p = doc.add_paragraph()
            for kw in data["scene_keywords"]:
                p.add_run(f"[{kw}] ")
            _set_para_font(p)

        doc.add_heading("3. 现场事故隐患总体评价", level=1)
        for run in doc.paragraphs[-1].runs:
            _set_run_font(run, "Microsoft YaHei", 12)
        p = doc.add_paragraph()
        p.add_run(f"重大事故隐患：{data['severe_count']} 项；").bold = True
        p.add_run(f"一般事故隐患：{data['general_count']} 项；").bold = True
        p.add_run(f"改进提升事项：{len(data['improvement_items'])} 项").bold = True
        _set_para_font(p)
        if data.get("overall_evaluation"):
            p2 = doc.add_paragraph(data["overall_evaluation"])
            _set_para_font(p2)

        severe_hazards = data.get("severe_hazards", [])
        general_hazards = data.get("general_hazards", [])
        # 处理 rowspan
        severe_hazards = self._add_rowspan(severe_hazards)
        general_hazards = self._add_rowspan(general_hazards)
        doc.add_heading("4. 现场事故隐患清单", level=1)
        for run in doc.paragraphs[-1].runs:
            _set_run_font(run, "Microsoft YaHei", 12)
        doc.add_heading("（一）重大事故隐患", level=2)
        for run in doc.paragraphs[-1].runs:
            _set_run_font(run, "Microsoft YaHei", 11)
        self._add_hazard_table(doc, severe_hazards, "隐患描述")
        doc.add_heading("（二）一般事故隐患", level=2)
        for run in doc.paragraphs[-1].runs:
            _set_run_font(run, "Microsoft YaHei", 11)
        self._add_hazard_table(doc, general_hazards, "隐患描述")

        improvement_items = data.get("improvement_items", [])
        improvement_items = self._add_rowspan(improvement_items)
        doc.add_heading("5. 现场改进提升事项清单", level=1)
        for run in doc.paragraphs[-1].runs:
            _set_run_font(run, "Microsoft YaHei", 12)
        self._add_hazard_table(doc, improvement_items, "改进事项")

        doc.add_heading("6. 下一步安全生产工作提醒", level=1)
        for run in doc.paragraphs[-1].runs:
            _set_run_font(run, "Microsoft YaHei", 12)
        for i, step in enumerate(data["next_steps_reminder"], 1):
            p3 = doc.add_paragraph(f"{i}. {step}")
            _set_para_font(p3)

        doc.add_paragraph()
        p1 = doc.add_paragraph("中国电建成都院智成安监系统")
        p1.runs[0].bold = True
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run_font(p1.runs[0])
        p2 = doc.add_paragraph(f"生成时间：{data['generated_time']}")
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run_font(p2.runs[0])

        return doc

    def generate_word_bytes(self, report_data: dict[str, Any]) -> bytes:
        doc = self._build_word_document(report_data)
        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def _add_hazard_table(self, doc: Any, hazards: list[dict[str, Any]], title: str = "隐患描述") -> None:
        """生成隐患表格，支持照片单元格纵向合并（rowspan）"""
        if not hazards:
            return

        # 按 photo_name 分组
        grouped: dict[str, list[dict[str, Any]]] = {}
        for item in hazards:
            photo = item.get("photo_name") or ""
            grouped.setdefault(photo, []).append(item)

        # 表头
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["序号", "现场照片", title, "判定依据", "备注"]
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = True
            _set_cell_font(cell)

        # 数据行
        seq = 1
        for photo_name, items in grouped.items():
            for idx, item in enumerate(items):
                row = table.add_row()
                row.cells[0].text = str(seq) if idx == 0 else ""
                if idx == 0:
                    cell = row.cells[1]
                    image_source = _get_image_source(item.get("photo_url"))
                    if image_source:
                        para = cell.paragraphs[0]
                        run = para.add_run()
                        run.add_picture(image_source, width=Cm(3.5))
                    else:
                        cell.text = "无"
                row.cells[2].text = item.get("description", "")
                row.cells[3].text = (
                    f"{item.get('standard_name', '')} {item.get('clause_no', '')}\n{item.get('clause_text', '')}"
                )
                row.cells[4].text = item.get("remark", "建议按现场实际复核")
                for cell in row.cells:
                    _set_cell_font(cell)
                seq += 1
        doc.add_paragraph()


def _build_cli() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="智成安监报告生成器")
    parser.add_argument(
        "--input-file",
        required=True,
        help="输入 JSON 文件路径，例如：output/report_20260519_101500.json",
    )
    parser.add_argument(
        "--format",
        choices=["pdf", "docx", "html", "all"],
        default="pdf",
        help="输出文件类型，默认 pdf",
    )
    parser.add_argument(
        "--name",
        default=None,
        help="输出文件名（不含扩展名）；未传时默认使用 uid",
    )
    parser.add_argument(
        "--template",
        default=None,
        help="可选 HTML 模板路径（默认使用内置模板）",
    )
    return parser


def _main() -> int:
    parser = _build_cli()
    args = parser.parse_args()
    config = _load_script_config()

    generator = ReportGenerator(template_path=args.template)
    try:
        data = json.loads(Path(args.input_file).read_text(encoding="utf-8"))
    except FileNotFoundError:
        parser.error(f"input-file 不存在: {args.input_file}")
    except OSError as exc:
        parser.error(f"读取 input-file 失败: {exc}")
    except json.JSONDecodeError as exc:
        parser.error(f"input-file 不是合法 JSON: {exc}")
    base_name = args.name or uuid4().hex

    input_json_path = Path(args.input_file)
    try:
        input_json_bytes = input_json_path.read_bytes()
    except Exception as exc:
        print(f"读取输入 JSON 失败: {exc}")
        return 1

    report_outputs: list[tuple[str, bytes, str]] = []
    try:
        if args.format in {"html", "all"}:
            report_outputs.append(
                (
                    f"{base_name}.html",
                    generator.generate_html_bytes(data),
                    "text/html; charset=utf-8",
                )
            )
        if args.format in {"pdf", "all"}:
            report_outputs.append(
                (
                    f"{base_name}.pdf",
                    generator.html_to_pdf_bytes(generator.generate_html(data)),
                    "application/pdf",
                )
            )
        if args.format in {"docx", "all"}:
            report_outputs.append(
                (
                    f"{base_name}.docx",
                    generator.generate_word_bytes(data),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            )
    except Exception as exc:
        print(f"生成失败: {exc}")
        return 1

    uploaded_outputs: dict[str, str] = {}
    uploader = ObjectStorageUploader(
        _config_env_or_default(config, "upload_target", DEFAULT_UPLOAD_TARGET),
        frontend_id=_config_env_or_default(config, "frontend_id", DEFAULT_FRONTEND_ID),
        user_id=_config_env_or_default(config, "user_id", DEFAULT_USER_ID),
    )
    try:
        uploaded_outputs = _upload_report_bundle(
            uploader=uploader,
            report_outputs=report_outputs,
            input_json_bytes=input_json_bytes,
            base_name=base_name,
        )
    except Exception as exc:
        print(f"上传失败: {exc}")
        return 1

    print(json.dumps(uploaded_outputs, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(_main())
